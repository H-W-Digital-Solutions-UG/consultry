from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUTPUT_DIR = ROOT / "output" / "doc"
INSPO_DIR = OUTPUT_DIR / "OfferTemplates_Inspo"
LOGO_PATH = Path("/Users/jules/Documents/H&W/hw logo white background.png")
DOCX_PATH = OUTPUT_DIR / "HW_Offer_Template.docx"
INDEX_PATH = INSPO_DIR / "OfferTemplates_Inspo_INDEX.md"
CREATED_DATE = "2026-07-01"


BRAND_BLACK = "111111"
BRAND_GRAY = "F2F4F5"
BRAND_LINE = "D8DEE2"
BRAND_ACCENT = "2F6F73"
TEXT_MUTED = "667085"


SOURCE_ROWS = [
    {
        "vendor": "Turntabl",
        "received": "2026-06-08 21:01 UTC",
        "message": "Fwd: RFP-T-BACK-0006 Provider-Agnostic Wallet Base Architecture - Turntabl",
        "sender": "Jose Junior / forwarded Turntabl offer",
        "files": "Turntabl_RFP-T-BACK-0006_Wallet_Base_Proposal.pdf",
        "files_short": "Turntabl wallet-base proposal PDF",
        "notes": "Use as a concise scope-to-delivery proposal reference.",
    },
    {
        "vendor": "Systango",
        "received": "2026-06-12 17:35 UTC",
        "message": "Re: RFP-T-BACK-0006 Provider-Agnostic Wallet Base Architecture - Systango",
        "sender": "Josh Hyatt <josh.hyatt@systango.com>",
        "files": "Systango_RFP-T-BACK-0006_Polity_Proposal.pdf",
        "files_short": "Systango Polity proposal PDF",
        "notes": "Use for offer framing, team structure, assumptions, and commercial layout ideas.",
    },
    {
        "vendor": "Rubicon",
        "received": "2026-06-09 15:59 UTC",
        "message": "Re: RFP-T-BACK-0006 Provider-Agnostic Wallet Base Architecture - Rubicon",
        "sender": "Pablo Soria <pablo.soria@rubicontech.io>",
        "files": (
            "Rubicon_RFP-T-BACK-0006_Technical_Analysis_Estimates.pdf; "
            "Rubicon_RFP-T-BACK-0006_Proposal_Implementation_Plan.pdf"
        ),
        "files_short": "Rubicon estimates PDF; Rubicon implementation plan PDF",
        "notes": "Use the separate analysis and implementation-plan split as an inspiration pattern.",
    },
]


STAGED_FILES = [
    {
        "name": "Turntabl_RFP-T-BACK-0006_Wallet_Base_Proposal.pdf",
        "source": "Polity_RFP-T-BACK-0006_Wallet_Base__Proposal_Turntabl[86169].pdf",
        "size": "307574 bytes",
    },
    {
        "name": "Systango_RFP-T-BACK-0006_Polity_Proposal.pdf",
        "source": "Systango_Polity_Proposal__RFP-T-BACK-0006_[86171].pdf",
        "size": "692965 bytes",
    },
    {
        "name": "Rubicon_RFP-T-BACK-0006_Technical_Analysis_Estimates.pdf",
        "source": (
            "Rubicon___Polity___RFP-T-BACK-0006_Provider-Agnostic_Wallet_Base_Architecture"
            "___Technical_Analysis___Estimates_-_Hojas_de_ca_lculo_de_Google[86170].pdf"
        ),
        "size": "121633 bytes",
    },
    {
        "name": "Rubicon_RFP-T-BACK-0006_Proposal_Implementation_Plan.pdf",
        "source": "Provider-Agnostic_Wallet_Base_Architecture___Proposal_Implementation_Plan[86173].pdf",
        "size": "3420967 bytes",
    },
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color=BRAND_LINE):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_text(cell, text, bold=False, color=BRAND_BLACK, size=9.5):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Aptos"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def format_table(table, header=True):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_border(cell)
            if header and row_index == 0:
                set_cell_shading(cell, BRAND_GRAY)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor.from_string(BRAND_BLACK)


def add_paragraph(document, text="", style=None, bold=False, italic=False, color=BRAND_BLACK):
    paragraph = document.add_paragraph(style=style)
    if text:
        run = paragraph.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.name = "Aptos"
        run.font.color.rgb = RGBColor.from_string(color)
    return paragraph


def add_placeholder(document, text):
    paragraph = add_paragraph(document)
    run = paragraph.add_run(text)
    run.italic = True
    run.font.name = "Aptos"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(TEXT_MUTED)
    return paragraph


def add_labeled_paragraph(document, label, body):
    paragraph = document.add_paragraph()
    label_run = paragraph.add_run(f"{label}: ")
    label_run.bold = True
    label_run.font.name = "Aptos"
    label_run.font.size = Pt(10.5)
    label_run.font.color.rgb = RGBColor.from_string(BRAND_BLACK)
    body_run = paragraph.add_run(body)
    body_run.font.name = "Aptos"
    body_run.font.size = Pt(10.5)
    body_run.font.color.rgb = RGBColor.from_string(BRAND_BLACK)
    return paragraph


def add_heading(document, text, level=1):
    paragraph = document.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "Aptos Display"
        run.font.color.rgb = RGBColor.from_string(BRAND_BLACK if level == 1 else BRAND_ACCENT)
    return paragraph


def add_bullets(document, items):
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        run = paragraph.add_run(item)
        run.font.name = "Aptos"
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor.from_string(BRAND_BLACK)


def add_simple_table(document, headers, rows, widths=None):
    table = document.add_table(rows=1, cols=len(headers))
    for col, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[col], header, bold=True)
    for row_values in rows:
        row = table.add_row()
        for col, value in enumerate(row_values):
            set_cell_text(row.cells[col], value)
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Inches(width)
    format_table(table)
    return table


def set_document_styles(document):
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(BRAND_BLACK)

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = styles[style_name]
        style.font.name = "Aptos Display"
        style.font.color.rgb = RGBColor.from_string(BRAND_BLACK)

    styles["Title"].font.name = "Aptos Display"
    styles["Title"].font.size = Pt(28)
    styles["Title"].font.bold = True
    styles["Title"].font.color.rgb = RGBColor.from_string(BRAND_BLACK)


def configure_section(section):
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)


def add_header_footer(document):
    section = document.sections[0]
    header = section.header.paragraphs[0]
    header.text = "H&W Digital Solutions | Offer Agreement"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.name = "Aptos"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(TEXT_MUTED)

    footer = section.footer.paragraphs[0]
    footer.text = "Confidential - draft for commercial and legal review before distribution"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.name = "Aptos"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(TEXT_MUTED)


def build_cover(document):
    logo_paragraph = document.add_paragraph()
    logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if LOGO_PATH.exists():
        logo_run = logo_paragraph.add_run()
        logo_run.add_picture(str(LOGO_PATH), width=Inches(1.7))

    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(48)

    title = document.add_paragraph(style="Title")
    title.add_run("H&W Offer Agreement")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    subtitle = document.add_paragraph()
    subtitle_run = subtitle.add_run("Statement of Work, Commercial Terms, and Conditions")
    subtitle_run.font.name = "Aptos Display"
    subtitle_run.font.size = Pt(18)
    subtitle_run.font.color.rgb = RGBColor.from_string(BRAND_ACCENT)

    note = document.add_paragraph()
    note_run = note.add_run(
        "Client-ready agreement baseline with practical delivery, acceptance, payment, change-control, confidentiality, IP, data, and warranty terms."
    )
    note_run.font.name = "Aptos"
    note_run.font.size = Pt(10.5)
    note_run.font.color.rgb = RGBColor.from_string(TEXT_MUTED)

    document.add_paragraph()
    metadata = [
        ("Client legal name", "[Client legal name]"),
        ("Opportunity / RFP reference", "[Reference]"),
        ("Offer reference", "[H&W offer reference]"),
        ("Version", "[Version]"),
        ("Date", "[YYYY-MM-DD]"),
        ("Prepared by", "H&W Digital Solutions"),
        ("Valid until", "[YYYY-MM-DD]"),
    ]
    table = document.add_table(rows=len(metadata), cols=2)
    for row_index, (label, value) in enumerate(metadata):
        set_cell_text(table.rows[row_index].cells[0], label, bold=True)
        set_cell_shading(table.rows[row_index].cells[0], BRAND_GRAY)
        set_cell_text(table.rows[row_index].cells[1], value, color=TEXT_MUTED)
    format_table(table, header=False)
    document.add_page_break()


def build_template_body(document):
    add_heading(document, "Document Control")
    add_simple_table(
        document,
        ["Version", "Date", "Author", "Change Summary"],
        [["1.0", "2026-07-01", "H&W Digital Solutions", "Agreement baseline with substantive terms and conditions"]],
        widths=[0.8, 1.2, 1.4, 4.6],
    )

    add_heading(document, "Agreement Snapshot")
    add_paragraph(
        document,
        "This Offer Agreement sets out the services H&W Digital Solutions will provide to the client, the acceptance and payment mechanics for those services, and the legal and operational terms that govern delivery.",
    )
    add_simple_table(
        document,
        ["Field", "Agreement Detail"],
        [
            ["Client", "[Client legal name and registration details]"],
            ["Project", "[Project name / opportunity / RFP reference]"],
            ["Commercial model", "[Fixed fee / time and materials / hybrid]"],
            ["Delivery term", "[Start date] to [target completion date], subject to dependencies and change control"],
            ["Primary contacts", "H&W: [name, email]; Client: [name, email]"],
            ["Offer validity", "This offer is valid until [YYYY-MM-DD] unless extended in writing by H&W."],
        ],
        widths=[2.1, 5.3],
    )

    add_heading(document, "1. Offer And Agreement Formation")
    add_paragraph(
        document,
        "1.1 This document becomes binding when signed by both parties or when the client issues a purchase order that expressly references this offer and H&W accepts it in writing."
    )
    add_paragraph(
        document,
        "1.2 If this offer conflicts with a master services agreement signed by both parties, the master services agreement prevails for legal terms and this offer prevails for project-specific scope, fees, milestones, and acceptance criteria."
    )
    add_paragraph(
        document,
        "1.3 Any client purchasing terms, portal terms, or invoice terms are excluded unless H&W expressly accepts them in a signed writing."
    )

    add_heading(document, "Scope Of Work")
    add_paragraph(
        document,
        "H&W will provide the professional services described in the work packages below. The scope should be read narrowly: only the listed work packages, deliverables, acceptance criteria, and expressly included activities are included."
    )
    add_simple_table(
        document,
        ["Workstream", "Included Activities", "Primary Output"],
        [
            ["Discovery and alignment", "Confirm objectives, constraints, stakeholders, inputs, acceptance path, and dependencies.", "Kickoff notes, clarified scope, and dependency register."],
            ["Implementation / advisory delivery", "Perform the agreed technical, product, advisory, or delivery activities for the project.", "Work products listed in the deliverables schedule."],
            ["Validation and handover", "Prepare evidence, conduct review sessions, resolve accepted defects, and hand over final materials.", "Acceptance evidence pack and final handover note."],
        ],
        widths=[1.8, 3.7, 2.1],
    )

    add_heading(document, "Out Of Scope", 2)
    add_bullets(
        document,
        [
            "Production activation, live credential use, regulated operations, or access to client secrets unless expressly included in the work package table.",
            "New integrations, environments, regulatory analysis, security certification, or third-party licensing not listed as deliverables.",
            "Ongoing managed service, support, monitoring, or incident response after final acceptance unless purchased under a support schedule.",
            "Any work triggered by changed client requirements, unavailable dependencies, undocumented systems, or third-party delays, except through change control.",
        ],
    )

    add_heading(document, "Deliverables And Acceptance Criteria")
    add_paragraph(
        document,
        "Each deliverable must have a clear acceptance gate. Acceptance is based on the stated criteria and evidence, not on unstated expectations."
    )
    add_simple_table(
        document,
        ["Deliverable", "Description", "Acceptance Criteria", "Evidence"],
        [
            ["D1 - Discovery pack", "Current-state findings, assumptions, and scope confirmation.", "Client confirms scope, dependencies, and unresolved blockers.", "Kickoff notes, decision log, dependency register."],
            ["D2 - Main delivery package", "Primary work products agreed for the project.", "Materially conforms to the agreed scope and work package criteria.", "Files, links, review notes, test results, or artefact index."],
            ["D3 - Final handover", "Closeout pack and transfer of usable deliverables.", "Accepted findings are resolved or logged as agreed residual risks.", "Handover note, acceptance memo, residual-risk list."],
        ],
        widths=[1.5, 2.3, 2.5, 1.3],
    )

    add_heading(document, "Acceptance Process", 2)
    add_bullets(
        document,
        [
            "H&W submits each deliverable with reasonable acceptance evidence, such as an artefact index, test output, review notes, screenshots, logs, or sign-off checklist.",
            "The client must accept or reject a submitted deliverable in writing within five business days. A rejection must identify the specific material non-conformity against the agreed criteria.",
            "If a deliverable is rejected, H&W will correct the accepted non-conformity and resubmit the deliverable. The review period restarts only for the corrected items.",
            "A deliverable is deemed accepted if the client uses it in production or materially relies on it for business purposes, or if no written rejection is received within the review period.",
            "Final Acceptance occurs when all deliverables are accepted, deemed accepted, or recorded as closed with agreed residual risks.",
        ],
    )

    document.add_page_break()
    add_heading(document, "Delivery Governance")
    add_paragraph(
        document,
        "The project will be run with lightweight governance. Decisions, blockers, scope questions, and acceptance evidence should be recorded in writing so that delivery and payment are not dependent on informal memory."
    )
    add_simple_table(
        document,
        ["Forum", "Cadence", "Participants", "Purpose"],
        [
            ["Delivery sync", "[Weekly / twice weekly]", "H&W lead and client delivery owner", "Progress, blockers, upcoming decisions, and evidence readiness."],
            ["Commercial checkpoint", "At each milestone or material scope question", "Commercial owners", "Fees, change requests, timeline impacts, and dependency pauses."],
            ["Acceptance review", "At deliverable submission", "Named reviewers", "Written acceptance, rejection, or residual-risk closure."],
        ],
        widths=[1.7, 1.4, 2.2, 2.3],
    )

    add_heading(document, "Dependencies, Assumptions, And Blocked Conditions")
    add_paragraph(
        document,
        "The fees, timeline, and acceptance path rely on the assumptions below. If a dependency is unavailable or materially different, H&W may pause affected work, record a blocked condition, and adjust delivery dates day-for-day until the dependency is resolved."
    )
    add_simple_table(
        document,
        ["Dependency / Assumption", "Owner", "If Unavailable Or Changed"],
        [
            ["Access to relevant documents, repositories, environments, data, and named reviewers.", "Client", "Affected work is paused or completed using agreed mocks, samples, or written assumptions."],
            ["Client answers clarifications and reviews submissions within agreed review periods.", "Client", "Milestone dates and payment triggers shift for the affected period."],
            ["Third-party tools, vendors, APIs, and licences behave as documented and remain available.", "Client / third party", "Impact is handled as a dependency issue or change request."],
            ["H&W can use remote delivery and AI-assisted engineering or drafting tools under the confidentiality controls below.", "H&W", "If restricted, fees and timeline may be revised before work continues."],
        ],
        widths=[3.2, 1.2, 3.0],
    )

    document.add_page_break()
    add_heading(document, "Commercial Terms")
    add_paragraph(
        document,
        "Unless the commercial schedule states otherwise, fees are exclusive of VAT, withholding tax, bank charges, and out-of-pocket expenses. Expenses require prior written approval and are billed at cost."
    )
    add_simple_table(
        document,
        ["Commercial Item", "Default Term"],
        [
            ["Fixed-fee work", "Fees are payable by milestone on written or deemed acceptance of the relevant milestone."],
            ["Time-and-materials work", "Fees are invoiced monthly in arrears based on time recorded at the agreed rates."],
            ["Payment term", "Invoices are payable within 14 calendar days of invoice date unless another period is stated in the commercial schedule."],
            ["Late payment", "H&W may suspend work on five business days' notice if undisputed invoices remain overdue."],
            ["Currency", "[EUR / USD / GBP], payable without deduction except taxes required by law."],
        ],
        widths=[2.0, 5.4],
    )

    add_heading(document, "Milestone Payment Model", 2)
    add_paragraph(
        document,
        "For fixed-fee projects, H&W may use the milestone model below. It balances startup funding, evidence-based progress, final acceptance, and a limited warranty holdback."
    )
    add_simple_table(
        document,
        ["Milestone", "Percentage", "Trigger"],
        [
            ["Kickoff", "10%", "Signed agreement or accepted purchase order, onboarding complete, and required access path confirmed."],
            ["Discovery acceptance", "10%", "Discovery pack accepted with no unresolved blocking findings."],
            ["Mid-term delivery", "20%", "Approximately 50% of agreed work packages delivered with evidence accepted."],
            ["Final QA submission", "30%", "Completed delivery package submitted with evidence and ready for client review."],
            ["Final Acceptance", "25%", "Client issues written Final Acceptance or all deliverables are deemed accepted."],
            ["Warranty release", "5%", "90 calendar days after Final Acceptance, subject to no open material warranty defect."],
        ],
        widths=[1.8, 1.0, 4.8],
    )

    add_heading(document, "Delay Treatment And Service Credits", 2)
    add_bullets(
        document,
        [
            "If a milestone is delayed solely by H&W and the commercial schedule includes service credits, the credit is 0.5% of the fixed fee per full week of delay, capped at 5% of the fixed fee.",
            "Service credits are applied against the next due invoice and are the client's exclusive financial remedy for non-material delay unless the delay is also a material breach.",
            "No service credit accrues for delays caused by client dependencies, review cycles, third-party systems, scope changes, force majeure, or agreed pauses.",
        ],
    )

    add_heading(document, "Change Control And Contingency")
    add_paragraph(
        document,
        "Any work outside the agreed scope requires a written change request approved before execution. H&W is not required to start changed work until commercial impact, timing impact, and acceptance criteria are agreed."
    )
    add_simple_table(
        document,
        ["Change Request Field", "Required Detail"],
        [
            ["Reason", "Why the change is required and who requested it."],
            ["Scope impact", "Added, removed, or amended deliverables and acceptance criteria."],
            ["Commercial impact", "Fixed fee adjustment, rate card impact, contingency use, or expense impact."],
            ["Timeline impact", "Milestones affected and any dependency pause or resequencing."],
            ["Approval", "Named approver for each party and date of written approval."],
        ],
        widths=[2.0, 5.4],
    )
    add_paragraph(
        document,
        "If a contingency allowance is included, it is not pre-approved spend. It may be used only after written approval that identifies the cause, amount, and affected deliverables."
    )

    add_heading(document, "Warranty, Defects, And Support")
    add_bullets(
        document,
        [
            "H&W warrants that it will perform the services with reasonable professional skill and care.",
            "For 30 calendar days after Final Acceptance, H&W will remediate material defects in accepted deliverables that are caused by H&W's failure to meet the agreed acceptance criteria.",
            "If the commercial schedule includes a 90-day warranty holdback, the warranty release is payable after the warranty period unless a material warranty defect is open.",
            "The warranty excludes defects caused by changed requirements, client modifications, third-party services, unsupported environments, misuse, or work outside scope.",
            "Post-acceptance support, monitoring, enhancements, or operational assistance must be purchased separately unless expressly included in the support schedule.",
        ],
    )

    add_heading(document, "Intellectual Property")
    add_bullets(
        document,
        [
            "Subject to full payment, the client owns the bespoke deliverables created specifically for the client under this agreement, excluding H&W background IP and third-party materials.",
            "H&W retains ownership of its pre-existing methods, templates, know-how, tools, libraries, reusable code, prompts, processes, and general skills developed or refined during delivery.",
            "To the extent H&W background IP is embedded in a deliverable, H&W grants the client a perpetual, worldwide, non-exclusive licence to use it as part of the deliverable for the client's internal business purposes.",
            "Open-source and third-party components remain governed by their own licence terms. H&W will not knowingly include copyleft code requiring disclosure of proprietary client code unless the client approves it in writing.",
            "H&W may reuse non-confidential learnings, patterns, and general know-how, provided it does not disclose the client's confidential information.",
        ],
    )

    add_heading(document, "Confidentiality And AI-Assisted Work")
    add_bullets(
        document,
        [
            "Each party must protect the other party's confidential information using at least reasonable care and may use it only to perform or receive the services.",
            "Confidentiality obligations last for five years after disclosure, except trade secrets, credentials, private keys, security vulnerabilities, and regulated data remain protected as long as they remain non-public.",
            "H&W may use AI-assisted drafting, coding, analysis, and review tools to improve speed and quality, but H&W remains responsible for human review of outputs before delivery.",
            "H&W will not submit client secrets, private keys, live credentials, raw production data, regulated personal data, unpublished security vulnerabilities, or client-confidential source code to public third-party AI tools without written client approval.",
            "If the client requires stricter AI tooling controls, those controls must be stated in the project schedule before work begins because they may affect cost and timeline.",
        ],
    )

    add_heading(document, "Data Protection And Security")
    add_bullets(
        document,
        [
            "If H&W processes personal data on behalf of the client, the parties will enter into a data processing agreement before processing begins.",
            "The client remains responsible for determining the lawfulness, accuracy, minimisation, and retention basis for personal data supplied to H&W.",
            "H&W will use least-privilege access, reasonable technical and organisational measures, and commercially reasonable safeguards for client systems and materials under H&W control.",
            "H&W will not request or store production secrets unless necessary for agreed work. Where secrets are required, they must be shared through an approved secret-management process, not email or chat.",
            "Each party will notify the other without undue delay after becoming aware of a security incident affecting the other party's confidential information or personal data.",
        ],
    )

    add_heading(document, "Client Responsibilities")
    add_bullets(
        document,
        [
            "Provide timely access to systems, documents, personnel, environments, licences, data, and third-party contacts needed for delivery.",
            "Ensure that supplied materials are accurate, complete, legally usable, and authorised for H&W to access and process.",
            "Review deliverables within the agreed acceptance windows and provide consolidated written feedback.",
            "Make business, legal, compliance, product, and architecture decisions that H&W reasonably needs to proceed.",
            "Pay undisputed invoices on time and raise invoice disputes in writing within seven calendar days of receipt.",
        ],
    )

    document.add_page_break()
    add_heading(document, "General Legal Terms")
    add_labeled_paragraph(
        document,
        "Liability cap",
        "Each party's aggregate liability is capped at the fees paid or payable under this agreement in the 12 months before the claim, unless the commercial schedule states another cap."
    )
    add_labeled_paragraph(
        document,
        "Excluded losses",
        "Neither party is liable for indirect loss, loss of profit, loss of revenue, loss of goodwill, or loss of anticipated savings, except where exclusion is prohibited by law."
    )
    add_labeled_paragraph(
        document,
        "Carve-outs",
        "The liability cap does not limit liability for fraud, wilful misconduct, unpaid fees, breach of confidentiality, infringement indemnities, or liability that cannot legally be limited."
    )
    add_labeled_paragraph(
        document,
        "Termination for convenience",
        "Either party may terminate for convenience on 10 business days' written notice. The client must pay for accepted work, work performed to date, and non-cancellable commitments."
    )
    add_labeled_paragraph(
        document,
        "Termination for cause",
        "Either party may terminate for material breach if the breach is not cured within 10 business days after written notice."
    )
    add_labeled_paragraph(
        document,
        "Suspension",
        "H&W may suspend work for overdue undisputed invoices, unsafe access requests, missing dependencies, or requests that would breach law, security obligations, or professional standards."
    )
    add_labeled_paragraph(
        document,
        "Subcontractors",
        "H&W may use employees, contractors, and specialist subcontractors, provided H&W remains responsible for their work and confidentiality obligations."
    )
    add_labeled_paragraph(
        document,
        "Force majeure",
        "Neither party is liable for delay or failure caused by events beyond reasonable control, provided it gives prompt notice and mitigates impact."
    )
    add_labeled_paragraph(
        document,
        "Non-solicitation",
        "During the project and for 12 months after final acceptance, neither party will knowingly solicit the other party's personnel involved in the project, except through general advertisements."
    )
    add_labeled_paragraph(
        document,
        "Governing law",
        "This agreement is governed by [governing law], and disputes are subject to the courts of [venue], unless mandatory law requires otherwise."
    )
    add_labeled_paragraph(
        document,
        "Notices",
        "Formal notices must be sent by email and registered post or courier to the addresses stated in the signature block or later notified in writing."
    )

    document.add_page_break()
    add_heading(document, "Schedules")
    add_heading(document, "Schedule 1 - Work Package Detail", 2)
    add_simple_table(
        document,
        ["ID", "Work Package", "Deliverable", "Acceptance Gate", "Fee / Effort"],
        [
            ["WP1", "[Name]", "[Deliverable]", "[Acceptance gate]", "[Amount / hours]"],
            ["WP2", "[Name]", "[Deliverable]", "[Acceptance gate]", "[Amount / hours]"],
            ["WP3", "[Name]", "[Deliverable]", "[Acceptance gate]", "[Amount / hours]"],
        ],
        widths=[0.5, 1.45, 2.0, 2.2, 1.3],
    )

    add_heading(document, "Schedule 2 - Fees And Milestones", 2)
    add_simple_table(
        document,
        ["Milestone", "Fee", "Invoice Trigger", "Acceptance Evidence"],
        [
            ["Kickoff", "[Amount]", "Signature / purchase order accepted", "Access and onboarding confirmation"],
            ["Discovery", "[Amount]", "Discovery pack submitted", "Decision log and dependency register"],
            ["Delivery", "[Amount]", "Delivery package submitted", "Artefact index and review evidence"],
            ["Final Acceptance", "[Amount]", "Final Acceptance achieved", "Written acceptance memo or deemed acceptance"],
            ["Warranty Release", "[Amount]", "Warranty period complete", "No open material warranty defect"],
        ],
        widths=[1.5, 1.1, 2.4, 2.4],
    )

    add_heading(document, "Schedule 3 - Optional Support", 2)
    add_simple_table(
        document,
        ["Support Item", "Duration", "Included", "Fee"],
        [
            ["Post-acceptance support", "[30 days / 1 month]", "Minor defect triage, handover questions, and agreed advisory support.", "[Amount / rate]"],
            ["Enhancement work", "As requested", "New features, expanded scope, or operational support.", "Change request or rate card"],
        ],
        widths=[1.7, 1.3, 3.2, 1.2],
    )

    add_heading(document, "Formal Agreement")
    add_paragraph(
        document,
        "By signing below, each party confirms that it has authority to enter into this agreement and agrees to be bound by the offer, schedules, and terms above."
    )
    add_simple_table(
        document,
        ["For H&W Digital Solutions", "For [Client Legal Name]"],
        [["Name: [Name]\nTitle: [Title]\nDate: [Date]\nSignature:", "Name: [Name]\nTitle: [Title]\nDate: [Date]\nSignature:"]],
        widths=[3.7, 3.7],
    )


def build_appendix(document):
    document.add_page_break()
    add_heading(document, "Appendix - Offer Inspiration Sources")
    add_paragraph(
        document,
        "The following recent Polity vendor offers were loaded into the local inspiration folder and used to shape this reusable template.",
    )
    rows = []
    for source in SOURCE_ROWS:
        rows.append(
            [
                source["vendor"],
                source["received"],
                source["message"],
                source["files_short"],
                source["notes"],
            ]
        )
    add_simple_table(
        document,
        ["Vendor", "Received", "Source Message", "Staged File(s)", "Template Signal"],
        rows,
        widths=[0.9, 1.2, 2.1, 2.0, 1.6],
    )


def write_index():
    lines = [
        "# OfferTemplates_Inspo Source Index",
        "",
        f"Created: {CREATED_DATE}",
        "",
        "## Purpose",
        "",
        "This folder contains recent Polity vendor offer examples plus the generated H&W offer template.",
        "The source emails were identified through the Outlook Email connector and the PDFs were copied from the local Outlook attachment cache with normalized filenames.",
        "",
        "## Generated Template",
        "",
        "- HW_Offer_Template.docx",
        "",
        "## Vendor Source Files",
        "",
    ]
    for item in STAGED_FILES:
        lines.extend(
            [
                f"- {item['name']}",
                f"  - Original cached filename: {item['source']}",
                f"  - Size: {item['size']}",
            ]
        )
    lines.extend(["", "## Outlook Source Messages", ""])
    for source in SOURCE_ROWS:
        lines.extend(
            [
                f"- {source['vendor']}",
                f"  - Received: {source['received']}",
                f"  - Sender: {source['sender']}",
                f"  - Subject: {source['message']}",
                f"  - Staged file(s): {source['files']}",
            ]
        )
    lines.extend(
        [
            "",
            "## Notes",
            "",
            "- The mailbox spelling for the vendor is Turntabl.",
            "- The Rubicon offer was supplied as a technical analysis / estimates PDF plus a separate proposal implementation plan PDF.",
            "- Do not distribute these inspiration files outside the approved working context without confirming vendor confidentiality terms.",
            "",
        ]
    )
    INDEX_PATH.write_text("\n".join(lines), encoding="utf-8")


def build_document():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    INSPO_DIR.mkdir(parents=True, exist_ok=True)

    document = Document()
    configure_section(document.sections[0])
    set_document_styles(document)
    add_header_footer(document)
    build_cover(document)
    build_template_body(document)
    document.save(DOCX_PATH)
    document.save(INSPO_DIR / DOCX_PATH.name)
    write_index()


if __name__ == "__main__":
    build_document()
