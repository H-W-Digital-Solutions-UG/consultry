from __future__ import annotations

from pathlib import Path
from textwrap import dedent

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parents[1]
OUTPUT_PATH = ROOT / "output" / "doc" / "Datenschutzrichtlinie-Consultry.docx"

POLICY_TEXT = dedent(
    """
    Datenschutzrichtlinie

    Gültig ab: 16-April-2026

    Aktualisiert am: 17-April-2026

    Diese Datenschutzerklärung erläutert die Richtlinien von https://consultry.de hinsichtlich der Erhebung, Verwendung, Offenlegung und dem Schutz personenbezogener Daten, die wir erheben, wenn Sie auf https://consultry.de (den „Dienst“) zugreifen. Diese Datenschutzerklärung definiert und beschreibt Ihre Datenschutzrechte sowie den Schutz, den Sie nach den geltenden Datenschutzgesetzen genießen.

    Durch die Nutzung unseres Dienstes stimmen Sie der Erhebung und Verwendung Ihrer personenbezogenen Daten gemäß dieser Datenschutzerklärung zu. Bitte greifen Sie nicht auf unseren Dienst zu und nutzen Sie ihn nicht, wenn Sie der Erhebung und Verwendung Ihrer Informationen, wie in dieser Datenschutzerklärung beschrieben, nicht zustimmen. Diese Datenschutzerklärung wurde mit Hilfe des CookieScript Generator für Datenschutzerklärungen erstellt.

    ## Welche Informationen wir erheben, zu welchen Zwecken und auf welcher Rechtsgrundlage

    Die nachstehenden Definitionen der Rechtsgrundlagen sind wie folgt zu verstehen:

    Berechtigtes Interesse: Interesse des Unternehmens oder eines Dritten, sofern Ihre Interessen oder Grundrechte und -freiheiten nicht überwiegen und deren Nutzung Ihrer personenbezogenen Daten in einem angemessenen Verhältnis zu Ihren Rechten und Freiheiten steht.

    Vertragserfüllung: Verarbeitung Ihrer personenbezogenen Daten, wenn dies für die Erfüllung eines Vertrags, dessen Vertragspartei Sie sind, oder zur Durchführung vorvertraglicher Maßnahmen erforderlich ist.

    Rechtliche Verpflichtungen: Verarbeitung Ihrer personenbezogenen Daten, wenn dies für die Erfüllung einer rechtlichen oder regulatorischen Verpflichtung erforderlich ist, der wir unterliegen.

    Einwilligung: Ihre Einwilligung bedeutet jede freiwillig erteilte, spezifische, informierte und unmissverständliche Willensbekundung, mit der Sie durch eine Erklärung oder eine eindeutige bestätigende Handlung Ihr Einverständnis mit der Verarbeitung personenbezogener Daten, die Sie betreffen, signalisieren. Wir können Ihre Einwilligung einholen, wenn wir keine andere Rechtsgrundlage für die Verarbeitung Ihrer Daten haben.

    Nutzungsdaten: Was wir erheben: Informationen darüber, wie Nutzer mit unserer Website oder Anwendung interagieren, einschließlich besuchter Seiten, Verweildauer auf Seiten, Klickpfade und Browserdetails. Zweck: Analyse des Nutzerverhaltens, Optimierung der Leistung und Verbesserung der Benutzererfahrung. Aufbewahrungsfrist: 365 Tage. Rechtsgrundlage: Legitimate Interests.

    Persönliche Identifikationsinformationen: Was wir erheben: E-Mail-Adresse, Telefonnummer (mobil). Zweck: Verwaltung von Benutzerkonten, Gewährleistung eines sicheren Zugangs und Bereitstellung personalisierter Inhalte. Aufbewahrungsfrist: 730 Tage. Rechtsgrundlage: Consent.

    Kontaktinformationen: Was wir erheben: Work Phone Number, Work E-Mail. Zweck: Sicherstellung einer korrekten Leistungserbringung und Unterstützung bei kontobezogenen Mitteilungen. Aufbewahrungsfrist: 730 Tage. Rechtsgrundlage: Consent.

    Online-Kennungen & Geräteinformationen: Was wir erheben: IP-Adresse, Browsertyp und -version, Betriebssystem. Zweck: Absicherung unserer Plattform, Optimierung der Leistung und Verständnis der technischen Nutzung. Aufbewahrungsfrist: 365 Tage. Rechtsgrundlage: Consent.

    Standortdaten: Was wir erheben: Geolocation (GPS-Koordinaten, Ort basierend auf IP-Adresse). Zweck: Personalisierung von Erlebnissen und Diensten basierend auf Ihrem Standort. Aufbewahrungsfrist: 365 Tage. Rechtsgrundlage: Consent.

    Kommunikationsdaten: Was wir erheben: E-Mail-Korrespondenz. Zweck: Bereitstellung von Kundensupport und Aufrechterhaltung einer Korrespondenzhistorie. Aufbewahrungsfrist: 730 Tage. Rechtsgrundlage: Legitimate Interests.

    Präferenz- und Verhaltensdaten: Was wir erheben: Interessen und Präferenzen der Nutzer, Interaktionsverlauf mit Inhalten oder Produkten. Zweck: Personalisierung von Empfehlungen, Werbung und Inhalten. Aufbewahrungsfrist: 730 Tage. Rechtsgrundlage: Consent.

    Berufliche Informationen: Was wir erheben: Firmenname, Geschäftliche E-Mail-Adresse, Team Size, Industry, Industry Focus. Zweck: Unterstützung von Geschäftsbeziehungen und Anpassung von Inhalten für berufliche Nutzer. Aufbewahrungsfrist: 730 Tage. Rechtsgrundlage: Consent.

    Einwilligungs- und Rechtsdokumente: Was wir erheben: Einwilligungsformulare. Zweck: Erfüllung rechtlicher und regulatorischer Verpflichtungen. Aufbewahrungsfrist: 730 Tage. Rechtsgrundlage: Consent.

    ## Wie wir Ihre personenbezogenen Daten erhalten

    Wir erheben Informationen, die Sie uns direkt bereitstellen, wenn Sie:

    - Formulare ausfüllen
    - Unsere Dienste nutzen
    - Mit uns korrespondieren
    - Uns aus anderen Gründen kontaktieren

    Through cookies, consent tools, and analytics technologies on our website, where required only with your consent.

    Wir können Ihre personenbezogenen Daten auch aus anderen legitimen Quellen beziehen:

    - Wir können personenbezogene Daten von Partnern und Kooperationsunternehmen erhalten

    ## Pflicht zur Bereitstellung personenbezogener Daten

    In bestimmten Fällen ist die Bereitstellung personenbezogener Daten gesetzlich oder vertraglich vorgeschrieben oder zur Durchführung eines Vertrags erforderlich.

    ## Wie wir Ihre Informationen teilen

    H&W Digital Solutions UG kann Ihre personenbezogenen Daten den Empfängern der folgenden Kategorien offenlegen:

    Mit Ihrer Einwilligung. H&W Digital Solutions UG wird Ihre Daten für jeden Zweck mit Ihrer ausdrücklichen Zustimmung teilen.

    Bei Unternehmensübertragungen. Ihre Daten werden im Falle einer Fusion, Übernahme, Verhandlung, eines Verkaufs von Geschäftsanteilen oder einer Finanzierung an Dritte übertragen.

    Wir übermitteln personenbezogene Daten in folgende Drittstaaten: Vereinigte Staaten. Dieses Land, Gebiet oder die angegebenen Sektoren wurden von der Europäischen Kommission als mit angemessenem Datenschutzniveau anerkannt.

    ## Konkret eingesetzte Dienstleister und Systeme

    Zur technischen Bereitstellung und zum Hosting unseres Dienstes nutzen wir derzeit insbesondere Vercel.

    Für das Consent-Management und die Steuerung von Einwilligungen verwenden wir CookieScript, einen Dienst von Objectis Ltd., Laisves st. 60, LT-05120 Vilnius, Litauen.

    Für optionale Statistik und Reichweitenmessung verwenden wir nach Ihrer Einwilligung Google Tag Manager und Google Analytics 4.

    Für Wartelisten-Anmeldungen, Double-Opt-In-Kommunikation und die Speicherung optionaler Angaben zur Lead-Qualifizierung verwenden wir Loops (loops.so), einen Dienst der Astrodon Corporation, USA.

    Soweit bei der Nutzung dieser Dienste personenbezogene Daten in Drittländer, insbesondere in die Vereinigten Staaten, übermittelt werden, erfolgt dies nach den jeweils von den Anbietern vorgesehenen Datenschutz- und Übermittlungsmechanismen.

    ## Rechte der betroffenen Person

    Sie haben folgende Rechte bezüglich Ihrer personenbezogenen Daten:

    - Recht auf Auskunft: Sie können eine Kopie der personenbezogenen Daten, die wir über Sie speichern, anfordern.
    - Recht auf Berichtigung: Sie können die Berichtigung unrichtiger oder unvollständiger Informationen verlangen.
    - Recht auf Löschung (Recht auf Vergessenwerden): Unter bestimmten Voraussetzungen können Sie die Löschung Ihrer personenbezogenen Daten verlangen.
    - Recht auf Einschränkung der Verarbeitung: Sie können die Einschränkung der Verarbeitung Ihrer personenbezogenen Daten unter bestimmten Umständen verlangen.
    - Recht auf Datenübertragbarkeit: Sie haben das Recht, Ihre personenbezogenen Daten in einem strukturierten, gängigen und maschinenlesbaren Format zu erhalten und diese Daten einem anderen Verantwortlichen zu übermitteln.
    - Widerspruchsrecht: In bestimmten Fällen können Sie der Verarbeitung Ihrer personenbezogenen Daten widersprechen.
    - Recht auf transparente Information: Sie haben das Recht, klare, transparente und leicht verständliche Informationen darüber zu erhalten, wie wir Ihre personenbezogenen Daten verarbeiten.

    Widerruf der Einwilligung: Gemäß Artikel 13 Absatz 2 Buchstabe c der DSGVO haben Sie das Recht, Ihre Einwilligung zur Verarbeitung Ihrer personenbezogenen Daten jederzeit zu widerrufen. Als Besucher dieser Website können Sie Ihre Einwilligung direkt über das auf der Website bereitgestellte Consent-Banner ganz einfach anpassen oder widerrufen. Alternativ können Sie sich auch an den Websitebetreiber wenden, dessen Kontaktdaten in dieser Datenschutzerklärung angegeben sind. Bitte beachten Sie, dass der Widerruf Ihrer Einwilligung die Rechtmäßigkeit der Verarbeitung, die vor dem Widerruf erfolgte, nicht berührt.

    Um diese Rechte auszuüben, kontaktieren Sie uns bitte unter contact@hw-digitalsolutions.de.

    Wir werden Ihre Rechte nur ausüben, nachdem wir Ihren schriftlichen Antrag zum Ausüben eines bestimmten Rechts erhalten und Ihre Identität verifiziert haben.

    Ihre Anträge werden innerhalb eines Monats ab dem Eingang des Antrags, der unseren internen Regelungen und der DSGVO entspricht, erfüllt oder eine Ablehnung mit Angabe der Gründe mitgeteilt. Diese Frist kann um zwei weitere Monate verlängert werden, wenn der Antrag aufgrund des Umfangs der personenbezogenen Daten oder weiterer gleichzeitig bearbeiteter Anträge besonders aufwendig ist. Wir werden Sie innerhalb eines Monats nach Eingang des Antrags über eine Verlängerung und die Gründe dafür informieren. Die Antwort erfolgt in der von Ihnen gewünschten Form.

    Wir können die Erfüllung Ihres Antrags verweigern, wenn die in der DSGVO vorgesehenen Ausnahmen und/oder Einschränkungen der Rechte der betroffenen Person Anwendung finden oder wenn Ihr Antrag offensichtlich unbegründet oder unverhältnismäßig ist. Im Falle einer Ablehnung werden wir Ihnen die Gründe dafür schriftlich mitteilen.

    ## Beschwerderecht bei einer Aufsichtsbehörde

    Sie haben das Recht, bei einer zuständigen Aufsichtsbehörde Beschwerde einzureichen, wenn Sie der Ansicht sind, dass die Verarbeitung Ihrer personenbezogenen Daten gegen geltendes Recht und Ihre Rechte und berechtigten Interessen verstößt. Sie können gemäß den für Bundesbeauftragte für den Datenschutz und die Informationsfreiheit (BfDI) der Bundesrepublik Deutschland festgelegten Verfahren vorgehen, die unter folgendem Link zu finden sind: https://www.bfdi.bund.de/anschriften.

    Bitte beachten Sie, dass trotz aller Bemühungen kein System vollständig sicher ist. Wir haben jedoch angemessene Sicherheitsmaßnahmen implementiert, um die Risiken unbefugten Zugriffs auf oder unsachgemäßen Verwendung Ihrer personenbezogenen Daten zu minimieren.

    ## Sicherheit

    Wir und unsere Drittanbieter, die im Auftrag der oben genannten Zwecke personenbezogene Daten verarbeiten können, sind vertraglich verpflichtet, die Vertraulichkeit dieser Daten zu wahren.

    Wir setzen angemessene Sicherheitspraktiken und -verfahren ein, um die Vertraulichkeit und Sicherheit Ihrer Informationen, einschließlich aller nicht öffentlichen personenbezogenen Daten, zu schützen.

    Wir schützen Ihre Daten durch angemessene physische, technische und administrative Sicherheitsmaßnahmen, einschließlich der Beschränkung des Zugriffs auf Ihre Informationen auf Mitarbeiter, die diese Informationen kennen müssen.

    ## Änderungen dieser Richtlinie

    Wir überprüfen diese Datenschutzerklärung regelmäßig und behalten uns das Recht vor, sie jederzeit in Übereinstimmung mit geltenden Gesetzen und Vorschriften zu ändern. Änderungen treten mit ihrer Veröffentlichung auf unserer Website sofort in Kraft.

    Bitte überprüfen Sie diese Datenschutzerklärung von Zeit zu Zeit, um über etwaige Änderungen informiert zu bleiben.

    ## Kontakt

    Für Fragen kontaktieren Sie uns bitte über folgende Methoden:

    Name: H&W Digital Solutions UG

    Adresse: Greifswalder Straße 13D, 10405 Berlin

    E-Mail: contact@hw-digitalsolutions.de

    Website: https://consultry.de

    Telefon: +4917670937993
    """
).strip()


def set_font(style, name: str, size: float, bold: bool = False) -> None:
    style.font.name = name
    style._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    style.font.size = Pt(size)
    style.font.bold = bold


def flush_paragraph(document: Document, paragraph_lines: list[str]) -> None:
    if not paragraph_lines:
        return
    document.add_paragraph(" ".join(paragraph_lines))
    paragraph_lines.clear()


def build_document() -> Document:
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2.3)

    styles = document.styles
    set_font(styles["Normal"], "Arial", 10.5)
    set_font(styles["Title"], "Arial", 18, bold=True)
    set_font(styles["Heading 1"], "Arial", 13.5, bold=True)

    lines = POLICY_TEXT.splitlines()
    title = ""
    body_started = False
    paragraph_lines: list[str] = []

    for raw_line in lines:
        line = raw_line.strip()
        if not body_started and line:
            title = line
            body_started = True
            break

    if not title:
        raise RuntimeError("No document title found.")

    document.add_paragraph(title, style="Title")

    title_consumed = False
    for raw_line in lines:
        line = raw_line.strip()

        if not line:
            flush_paragraph(document, paragraph_lines)
            continue

        if not title_consumed:
            if line == title:
                title_consumed = True
            continue

        if line.startswith("## "):
            flush_paragraph(document, paragraph_lines)
            document.add_heading(line[3:], level=1)
            continue

        if line.startswith("- "):
            flush_paragraph(document, paragraph_lines)
            document.add_paragraph(line[2:], style="List Bullet")
            continue

        paragraph_lines.append(line)

    flush_paragraph(document, paragraph_lines)
    return document


def main() -> None:
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
