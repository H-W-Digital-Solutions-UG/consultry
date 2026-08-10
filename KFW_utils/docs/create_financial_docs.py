from pathlib import Path
import sys

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


GREEN = "0F5B4D"
DARK = "17332D"
MINT = "DCEBE6"
PALE = "F2F5F3"
WARM = "FFF0DB"
RED = "9C2F24"
GRAY = "59645F"
WHITE = "FFFFFF"
LOGO = Path("design/logos/consultry-weblogo-cropped.png").resolve()
OUT_DOC = Path("output/doc")


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell, top=55, start=75, bottom=55, end=75):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def no_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])


def bottom_border(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "5")
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), GREEN)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def setup_doc(title, short_title, landscape=False):
    doc = Document()
    section = doc.sections[0]
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Cm(29.7)
        section.page_height = Cm(21.0)
        section.left_margin = Cm(1.2)
        section.right_margin = Cm(1.2)
        section.top_margin = Cm(1.4)
        section.bottom_margin = Cm(1.25)
    else:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(1.7)
        section.right_margin = Cm(1.6)
        section.top_margin = Cm(1.6)
        section.bottom_margin = Cm(1.5)
    section.header_distance = Cm(0.55)
    section.footer_distance = Cm(0.55)

    header = section.header.paragraphs[0]
    try:
        header.add_run().add_picture(str(LOGO), width=Inches(0.75))
    except Exception:
        r = header.add_run("CONSULTRY")
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(GREEN)
    r = header.add_run(f"    {short_title} | KfW StartGeld 067 | ENTWURF")
    r.font.name = "Aptos"
    r.font.size = Pt(7.5)
    r.font.color.rgb = RGBColor.from_string(GRAY)
    bottom_border(header)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run("Consultry Software GmbH | 12.07.2026 | Seite ")
    r.font.name = "Aptos"
    r.font.size = Pt(7)
    r.font.color.rgb = RGBColor.from_string(GRAY)
    page_field(footer)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(8.8 if landscape else 9.2)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.04

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Aptos Display"
    h1.font.size = Pt(21)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor.from_string(GREEN)
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after = Pt(5)

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Aptos Display"
    h2.font.size = Pt(12)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor.from_string(DARK)
    h2.paragraph_format.space_before = Pt(7)
    h2.paragraph_format.space_after = Pt(3)

    doc.core_properties.title = title
    doc.core_properties.subject = "ERP-Gründerkredit StartGeld 067"
    doc.core_properties.author = "Consultry Software GmbH"
    return doc


def title_block(doc, title, subtitle):
    p = doc.add_paragraph(style="Heading 1")
    p.add_run(title)
    p = doc.add_paragraph()
    r = p.add_run(subtitle)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(GRAY)
    r.font.size = Pt(10)

    box = doc.add_table(rows=1, cols=1)
    box.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = box.cell(0, 0)
    shade(cell, WARM)
    margins(cell, 80, 110, 80, 110)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(
        "ENTWURF. Alle nicht belegten Zahlen sind Planannahmen. Vor Bankeinreichung sind Registerdaten, "
        "Geschäftsführergehälter, Steuerbehandlung, Umsatzsteuer, Pilot- und Context-Activation-Verträge zu bestätigen."
    )
    r.bold = True
    r.font.name = "Aptos"
    r.font.size = Pt(8.2)
    r.font.color.rgb = RGBColor.from_string(DARK)
    doc.add_paragraph()


def make_table(doc, headers, rows, widths=None, font_size=8.0, total_rows=None, open_cells=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False if widths else True
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        c = table.rows[0].cells[i]
        c.text = str(header)
        shade(c, GREEN)
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        margins(c)
        if widths:
            c.width = Cm(widths[i])
        for p in c.paragraphs:
            p.paragraph_format.space_after = Pt(0)
            for r in p.runs:
                r.font.name = "Aptos"
                r.font.size = Pt(font_size)
                r.bold = True
                r.font.color.rgb = RGBColor.from_string(WHITE)
    repeat_header(table.rows[0])
    total_rows = set(total_rows or [])
    open_cells = set(open_cells or [])
    for r_i, row_data in enumerate(rows, start=1):
        cells = table.add_row().cells
        no_split(table.rows[-1])
        for c_i, value in enumerate(row_data):
            c = cells[c_i]
            c.text = str(value)
            c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            margins(c)
            if widths:
                c.width = Cm(widths[c_i])
            if r_i in total_rows:
                shade(c, MINT)
            elif r_i % 2 == 0:
                shade(c, PALE)
            if (r_i, c_i) in open_cells:
                shade(c, WARM)
            for p in c.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                if c_i > 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for r in p.runs:
                    r.font.name = "Aptos"
                    r.font.size = Pt(font_size)
                    if r_i in total_rows:
                        r.bold = True
                    if (r_i, c_i) in open_cells:
                        r.bold = True
                        r.font.color.rgb = RGBColor.from_string(RED)
    return table


def add_note(doc, title, text, fill=PALE):
    box = doc.add_table(rows=1, cols=1)
    no_split(box.rows[0])
    cell = box.cell(0, 0)
    shade(cell, fill)
    margins(cell, 90, 110, 90, 110)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title + "\n")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(GREEN)
    r.font.size = Pt(9)
    r = p.add_run(text)
    r.font.size = Pt(8.3)
    r.font.color.rgb = RGBColor.from_string(DARK)
    return box


def fmt(x, decimals=1):
    if isinstance(x, str):
        return x
    return f"{x:,.{decimals}f}".replace(",", "X").replace(".", ",").replace("X", ".")


def create_rentability():
    doc = setup_doc(
        "Consultry Rentabilitätsvorschau",
        "Rentabilitätsvorschau",
        landscape=False,
    )
    title_block(doc, "Rentabilitätsvorschau", "Geschäftsjahre 1 bis 3 | Beträge in TEUR, netto, ohne AfA")

    doc.add_paragraph("Umsatztreiber", style="Heading 2")
    driver_rows = [
        ("zahlende Endkunden zum Jahresende (exkl. Design-Partner)", "4", "12", "24"),
        ("Design-Partner zum API-Token-Selbstkostenpreis", "2", "2", "2"),
        ("Preis je aktivem Seat / Monat", "50 EUR Pilot", "ab 69 EUR", "ab 69 EUR"),
        ("bezahlte Seats je kommerziellem Endkunden", "45", "45", "45"),
        ("End-ARR", "108,0", "406,1", "853,2"),
        ("SaaS-/Pilotumsatz im Geschäftsjahr", "36,0", "257,0", "614,1"),
        ("Context-Activation-Pakete", "20,0", "40,0", "60,0"),
        ("realisierter Nettoumsatz gesamt", "56,0", "297,0", "674,1"),
    ]
    make_table(
        doc,
        ["Treiber", "Jahr 1", "Jahr 2", "Jahr 3"],
        driver_rows,
        widths=[10.2, 2.6, 2.6, 2.6],
        font_size=7.25,
        total_rows={5, 8},
    )
    doc.add_paragraph()
    add_note(
        doc,
        "Herleitung",
        "Jahr 1 enthält vier zahlende Kunden ab M6/M8/M10/M12 mit 45 Seats zu 50 EUR pro Monat sowie vier Context-"
        "Activation-Pakete zu je 5 TEUR. Die zwei vorhandenen Design-Partner-LOIs laufen ausschließlich zum API-Token-"
        "Selbstkostenpreis und sind nicht als Umsatz enthalten. Ab Jahr 2 gilt 69 EUR nur für neue kommerzielle Kohorten. "
        "Jahresvorauszahlungen, Win-Fees und unbewiesene Reuse-/Erfolgsumsätze sind nicht angesetzt.",
    )

    doc.add_paragraph("Personalkosten Jahr 1", style="Heading 2")
    salary_rows = [
        ("CEO / Commercial", "0,0", "0,0", "0,0", "0,0"),
        ("CTO / Security & Engineering", "0,0", "0,0", "0,0", "0,0"),
        ("CPO / Product", "0,0", "0,0", "0,0", "0,0"),
        ("Verbindliche Personalkosten im Base-Case", "0,0", "0,0", "0,0", "0,0"),
        ("Konditionaler Personalrahmen nach Zusatzförderung", "-", "-", "max. 80,0", "nicht im Base-Case"),
    ]
    make_table(
        doc,
        ["Rolle", "Brutto M1-3", "Brutto M4-6", "Brutto M7-12", "AG-Kosten Jahr 1"],
        salary_rows,
        widths=[7.4, 2.6, 2.6, 2.6, 3.0],
        font_size=7.0,
        total_rows={4, 5},
    )

    doc.add_paragraph("Ergebnisvorschau", style="Heading 2")

    rows = [
        ("1. Nettoumsatz", "56,0", "297,0", "674,1"),
        ("2. Direkter Wareneinsatz / variable Cloud- und LLM-COGS", "4,3", "30,8", "73,7"),
        ("3. Rohertrag / Rohgewinn", "51,7", "266,2", "600,4"),
        ("4.1 Personalkosten inkl. Lohnnebenkosten", "0,0", "198,6", "351,0"),
        ("4.1.1 davon Gründer-/Geschäftsführergehälter inkl. AG-Anteil", "0,0", "165,6", "165,6"),
        ("4.2 Miete / Coworking", "3,0", "8,0", "12,0"),
        ("4.3 Heizung, Strom, Wasser", "0,5", "1,0", "2,0"),
        ("4.4 Werbung", "6,0", "12,0", "25,0"),
        ("4.5 Kraftfahrzeugkosten", "0,0", "0,0", "0,0"),
        ("4.6 Reisekosten", "4,0", "8,0", "15,0"),
        ("4.7 Telefon, Internet", "1,5", "3,0", "5,0"),
        ("4.8 Büromaterial", "0,5", "1,0", "2,0"),
        ("4.9 Verpackung", "0,0", "0,0", "0,0"),
        ("4.10 Reparatur / Instandhaltung", "0,5", "1,0", "2,0"),
        ("4.11 Versicherungen", "2,5", "4,0", "6,0"),
        ("4.12 Beiträge / Kammern", "0,5", "1,0", "2,0"),
        ("4.13 Leasing", "0,0", "0,0", "0,0"),
        ("4.14 Buchführung / Beratung", "6,0", "10,0", "15,0"),
        ("4.15 Sonstige Aufwendungen inkl. fixer Tech-Stack", "13,0", "11,0", "19,0"),
        ("Summe Aufwendungen 4.1 bis 4.15", "38,0", "258,6", "456,0"),
        ("5. Ergebnis vor Steuern, AfA und Zinsen", "13,7", "7,6", "144,4"),
        ("6. Zinsen", "9,0", "9,0", "8,5"),
        ("7. Ertragsteuern - vorläufig nach Verlustvortrag", "0,0", "0,0", "0,0"),
        ("8. Jahresüberschuss / -defizit vor AfA", "4,6", "-1,4", "135,9"),
        ("9. Tilgung - nur Liquidität", "0,0", "0,0", "25,0"),
        ("10. Private Versicherungen - GmbH nicht angesetzt", "0,0", "0,0", "0,0"),
        ("11. Verfügbarer Betrag nach Tilgung", "4,6", "-1,4", "110,9"),
    ]
    make_table(
        doc,
        ["Position", "Jahr 1", "Jahr 2", "Jahr 3"],
        rows,
        widths=[10.2, 2.6, 2.6, 2.6],
        font_size=7.35,
        total_rows={3, 20, 21, 24, 27},
    )

    doc.add_paragraph()
    add_note(
        doc,
        "Planungslogik",
        "End-ARR: 108,0 / 406,1 / 853,2 TEUR. Zahlende Kunden am Jahresende: 4 / 12 / 24, jeweils zusätzlich zu den "
        "zwei Token-Selbstkosten-Design-Partnern. Der Umsatz entsteht aus expliziten Startmonaten, Seats und Preisen. "
        "Wareneinsatz umfasst 12 Prozent der wiederkehrenden Umsätze; fixe Cloud-/Tech-Kosten stehen unter sonstigen Aufwendungen. "
        "Die KfW-Vorlage weist das Ergebnis ohne AfA aus.",
    )
    doc.add_paragraph()
    add_note(
        doc,
        "Personal- und Kosten-Gates",
        "Jahr 1 enthält keine festen Gründergehälter und keine verbindlichen Personalkosten. Ein Personalrahmen bis maximal "
        "80 TEUR darf erst nach gesonderter zusätzlicher Förderung und Gesellschafterfreigabe ausgelöst werden; er ist nicht "
        "im Base-Case finanziert. Ab Jahr 2 werden Gründergehälter angesetzt; erster Engineer ab M20, Customer Success ab M28 "
        "und Sales ab M31. Personalkosten im Base-Case: 0,0 / 198,6 / 351,0 TEUR.",
    )
    doc.add_paragraph()
    add_note(
        doc,
        "Vor Einreichung",
        "Vor Einreichung sind private Lebensunterhaltssicherung, zeitlicher Gründerumfang, Sozialversicherung sowie das Fehlen "
        "nachträglich fälliger oder gestundeter Gehaltsansprüche zu belegen. Das positive Jahr-1-EBITDA ist nicht normalisiert, "
        "sondern beruht auf unentgeltlicher Gründerarbeit. Steuerberater muss Verlustvortrag, "
        "Körperschaftsteuer, Gewerbesteuer, Aktivierung und Abschreibung der Software sowie Umsatzsteuer bestätigen. "
        "Der Kapitaldienst in Jahr 3 beträgt bei 10/2/10 und 4,52 Prozent rund 33,5 TEUR; vereinfachter EBITDA-DSCR ca. 4,3x. "
        "Cash-Steuern sind vorläufig mit null angesetzt, weil die steuerliche Abschreibung/Aktivierung des Softwarepakets das "
        "niedrige Jahr-1-Ergebnis vor AfA voraussichtlich überlagert; Steuerberater muss dies bestätigen.",
        fill=WARM,
    )
    doc.add_paragraph()
    add_note(
        doc,
        "Quellenbasis",
        "KfW-Formular 600 000 5284; KfW-Merkblatt 067, Stand 18.06.2026; KfW-Produktseite abgerufen 12.07.2026; "
        "Consultry Product Vision v2.7, MVP-PRD, GTM Decisions und Onboarding-/Korpus-Ritual. Die OmniSEC-Arbeitsmappe wurde nur "
        "als strukturelle Inspiration für Treiber-, Rollen- und Szenariologik verwendet; keine OmniSEC-Zahl wurde übernommen.",
    )

    path = OUT_DOC / "02_Consultry_Rentabilitaetsvorschau.docx"
    doc.save(path)
    return path


def create_liquidity():
    doc = setup_doc(
        "Consultry Liquiditätsplan",
        "Liquiditätsplan",
        landscape=True,
    )
    title_block(doc, "Liquiditätsplan", "Monate 1 bis 12 | Beträge in TEUR | Planungsansicht netto")

    months = [f"M{i}" for i in range(1, 13)]
    saas = [0, 0, 0, 0, 0, 0, 2.25, 2.25, 4.5, 4.5, 6.75, 6.75]
    activation = [0, 0, 0, 0, 0, 5, 0, 5, 0, 5, 0, 5]
    sales = [saas[i] + activation[i] for i in range(12)]
    loan = [200] + [0] * 11
    equity = [0] * 12
    investment = [26, 0, 20, 4, 20, 0, 20, 0, 15, 0, 15, 0]
    personnel = [0] * 12
    fixed_opex = [3.1666666667] * 12
    direct_cogs = [0, 0, 0, 0, 0, 0.27, 0.27, 0.54, 0.54, 0.81, 0.81, 1.08]
    interest = [0.7533333333] * 12

    inflow = [sales[i] + loan[i] + equity[i] for i in range(12)]
    outflow = [
        investment[i] + personnel[i] + fixed_opex[i] + direct_cogs[i] + interest[i]
        for i in range(12)
    ]
    saldo = [inflow[i] - outflow[i] for i in range(12)]
    cumulative = []
    running = 0
    for value in saldo:
        running += value
        cumulative.append(running)

    def row(label, values, decimals=1):
        return [label] + [fmt(v, decimals) for v in values] + [fmt(sum(values), decimals)]

    rows = [
        row("1.1 SaaS-Zahlungseingänge (1 Monat Lag)", saas, 2),
        row("1.2 Context-Activation-Zahlungseingänge", activation, 2),
        row("1.3 KfW-Darlehen", loan),
        row("1.4 Eigenmittel / Stammkapital", equity),
        row("1.5 Summe Liquiditätszugang", inflow),
        row("2.1 Anlageinvestitionen", investment),
        row("2.2 Personal", personnel),
        row("2.3 fixe Betriebsausgaben", fixed_opex, 3),
        row("2.4 variable Cloud-/LLM-COGS", direct_cogs, 3),
        row("2.5 sonstige Auszahlungen", [0] * 12),
        row("2.6 Kredittilgung", [0] * 12),
        row("2.7 Zinsen", interest, 3),
        row("2.8 Summe Liquiditätsabgang", outflow, 3),
        row("3. Monatlicher Liquiditätssaldo", saldo, 3),
        ["4. Liquidität kumuliert"] + [fmt(v, 3) for v in cumulative] + [fmt(cumulative[-1], 3)],
    ]

    headers = ["Position"] + months + ["Summe"]
    make_table(
        doc,
        headers,
        rows,
        widths=[4.5] + [1.55] * 13,
        font_size=6.45,
        total_rows={5, 13, 14, 15},
    )
    doc.add_paragraph()
    add_note(
        doc,
        "Ergebnis",
        "Endliquidität Jahr 1: 75,640 TEUR. Der niedrigste Cashbestand im 36-Monats-Base-Case liegt bei 41,046 TEUR in "
        "Jahr 2 und damit über dem internen 30-TEUR-Puffer. Der Base-Case unterstellt keine zusätzlichen Eigenmittel und "
        "keine festen Gründergehälter in Jahr 1. Vier zahlende Kunden starten in M6/M8/M10/M12; SaaS-Zahlungen laufen mit einem "
        "Monat Verzögerung. Die zwei Token-Selbstkosten-Design-Partner sind nicht als Umsatz enthalten.",
    )
    doc.add_paragraph()
    add_note(
        doc,
        "Steuer- und Einreichungshinweis",
        "Die KfW-Vorlage verlangt Umsätze inklusive Umsatzsteuer sowie Steuerzahlungen. Diese Planungsansicht ist netto. "
        "Vor Einreichung müssen deutsche Umsatzsteuer, B2B-Reverse-Charge in Österreich, Vorsteuer, Zahlungsziele, "
        "Körperschaft-/Gewerbesteuer und der tatsächliche Darlehensabruf monatlich durch den Steuerberater eingearbeitet werden.",
        fill=WARM,
    )
    doc.add_paragraph()
    add_note(
        doc,
        "Liquiditätssteuerung",
        "Zusätzlich wird rollierend eine 13-Wochen-Vorschau geführt. Bei erwarteter Liquidität unter 30 TEUR innerhalb "
        "von 90 Tagen gelten Spend-Freeze, Hiring-Stopp und Gesellschafterentscheidung. Der konditionale Personalrahmen von "
        "maximal 80 TEUR darf erst nach zusätzlicher Förderung freigegeben werden. Ohne Zusatzmittel würde seine volle "
        "Ausschöpfung den 36-Monats-Tiefpunkt auf rund -38,954 TEUR senken.",
    )

    doc.add_paragraph("Szenario-Test Jahr 1", style="Heading 2")
    downside_rows = [
        ("Base: 4 zahlende Kunden / 4 Activations", "75,64", "0 TEUR Zusatz-Eigenmittel; kein festes Gründergehalt"),
        ("Downside: 50 % kommerzieller Cash, 15 % Fixkosten-Gate", "60,00", "kein festes Gründergehalt"),
        ("Nur LOIs: kein Paid-Customer-Cash, 20 % Fixkosten-Gate", "40,56", "Design-Partner validieren Nutzung, nicht Preis"),
        ("Mindestliquidität 36-Monats-Base", "41,05", "Tiefpunkt Jahr 2; 30-TEUR-Management-Puffer"),
        ("Volle 80 TEUR Personal ohne Zusatzförderung", "-4,36", "Jahr-1-Ende; 36-Monats-Tiefpunkt ca. -38,95"),
        ("Zusatzliquidität für 80-TEUR-Freigabe", "68,95", "erforderlich für 30-TEUR-Puffer über 36 Monate"),
    ]
    make_table(
        doc,
        ["Position", "TEUR", "Annahme / Maßnahme"],
        downside_rows,
        widths=[8.0, 3.0, 12.0],
        font_size=7.5,
        total_rows={4, 5, 6},
    )
    doc.add_paragraph()
    add_note(
        doc,
        "Voraussetzung des Downside-Tests",
        "In Jahr 1 bestehen keine festen Gründergehälter und keine gestundeten Gehaltsansprüche im Base-Case. Der optionale "
        "Personalrahmen ist eine Obergrenze nach Zusatzförderung, kein bereits finanziertes Budget. Jahresvorauszahlungen, "
        "zusätzliche bezahlte Kunden und Design-Partner-Umsätze sind nicht zur Schließung unterstellt.",
        fill=WARM,
    )
    doc.add_paragraph()
    add_note(
        doc,
        "Quellenbasis",
        "KfW-Formular 600 000 5283; KfW-Merkblatt 067, Stand 18.06.2026; KfW-Produktseite abgerufen 12.07.2026; "
        "Consultry Businessplan, GTM Decisions und Onboarding-/Korpus-Ritual. OmniSEC diente ausschließlich als Inspiration für Burn-, Cash- "
        "und Downside-Darstellung.",
    )

    path = OUT_DOC / "03_Consultry_Liquiditaetsplan.docx"
    doc.save(path)
    return path


def create_investment():
    doc = setup_doc(
        "Consultry Investitionsplan",
        "Investitionsplan",
        landscape=False,
    )
    title_block(doc, "Investitionsplan", "Kapitalbedarf und Finanzierung | Beträge netto")

    doc.add_paragraph(
        "Der modellierte Finanzierungsbedarf des verbindlichen Base-Case beträgt 200.000 EUR und wird über den "
        "ERP-Gründerkredit - StartGeld finanziert: 120.000 EUR Investitionen und höchstens 80.000 EUR Betriebsmittel. "
        "In Jahr 1 bestehen keine festen Gründergehälter. Ein konditionaler Personalrahmen bis 80.000 EUR ist nicht Teil "
        "dieser Finanzierung und darf erst nach gesonderter zusätzlicher Förderung freigegeben werden."
    )

    doc.add_paragraph("A. Investitionen", style="Heading 2")
    investment_rows = [
        ("Externes, abnahmefähiges MVP-Software-/IP-Paket", "110.000 EUR", "Angebote, Pflichtenheft, IP-Übertragung, Aktivierung, Abnahme"),
        ("Entwicklungs-, Test- und Security-Hardware", "10.000 EUR", "Angebote, Inventarisierung"),
        ("Zwischensumme Investitionen", "120.000 EUR", "bankseitige Förderfähigkeitsbestätigung erforderlich"),
    ]
    make_table(
        doc,
        ["Position", "Betrag", "Nachweis"],
        investment_rows,
        widths=[8.2, 3.1, 6.7],
        font_size=7.8,
        total_rows={3},
    )

    doc.add_paragraph("A.1 Arbeitsstruktur des externen Software-/IP-Pakets", style="Heading 2")
    package_rows = [
        ("Source-bound Context Graph, Korpus- und Consulting-Objektmodell", "28.000 EUR"),
        ("Consultry Engine, HarnessPack-Vertrag, Authentisierung und Security", "25.000 EUR"),
        ("Opportunity-to-Concept, Evidence Pack und Grounding-Pipeline", "25.000 EUR"),
        ("Rollen-Workspace, anonyme TeamShape und Human Approval", "16.000 EUR"),
        ("Tests, technische Dokumentation, Deployment und Rechteübergabe", "16.000 EUR"),
        ("Summe externes Software-/IP-Paket", "110.000 EUR"),
    ]
    make_table(
        doc,
        ["Abnahmefähiges Arbeitspaket", "Betrag"],
        package_rows,
        widths=[14.9, 3.1],
        font_size=7.55,
        total_rows={6},
    )
    doc.add_paragraph()
    add_note(
        doc,
        "Beschaffungs- und Abnahmeregel",
        "Der Dienstleistervertrag muss Fest-/Meilensteinpreise, objektive Abnahmekriterien, vollständige Quellcode-, "
        "Dokumentations- und Nutzungsrechtsübertragung sowie Zahlung nach dokumentierter Teilabnahme enthalten. "
        "Interne Gründerarbeit ist nicht Bestandteil der 110.000-EUR-Investitionsposition. Das graphfähige Fundament darf "
        "spätere Project-Symbiosis-/Assetization-Flows ermöglichen; deren vollständige Realisierung ist nicht Bestandteil des H1-Pakets.",
    )

    doc.add_paragraph("B. Betriebsmittel- und Liquiditätsbedarf", style="Heading 2")
    operating_rows = [
        ("LLM-API, EU-Cloud und technische SaaS-Werkzeuge", "24.000 EUR", "Betriebsmittel"),
        ("Design-Partner-Onboarding, Vertrieb, Marketing und Reisen", "20.000 EUR", "Betriebsmittel"),
        ("Recht, Datenschutz, Buchführung und Versicherungen", "12.000 EUR", "Betriebsmittel"),
        ("Arbeitsplatz, Kommunikation und allgemeine Tools", "6.000 EUR", "Betriebsmittel"),
        ("externe Aktivierungs-/Datenvorbereitung und Support", "10.000 EUR", "Betriebsmittel"),
        ("Liquiditäts-/Unvorhergesehenes-Puffer", "8.000 EUR", "Betriebsmittel"),
        ("KfW-finanzierter Betriebsmittelanteil", "80.000 EUR", "Programmobergrenze"),
        ("davon feste Gründer-/Personalkosten Jahr 1", "0 EUR", "keine Verpflichtung oder Gehaltsstundung"),
    ]
    make_table(
        doc,
        ["Position", "Betrag", "Einordnung"],
        operating_rows,
        widths=[8.2, 3.1, 6.7],
        font_size=7.8,
        total_rows={7, 8},
    )

    doc.add_paragraph("C. Gesamtbedarf und Finanzierung", style="Heading 2")
    funding_rows = [
        ("Gesamter Finanzierungsbedarf Base-Case", "200.000 EUR", "120.000 Investition + 80.000 Betriebsmittel"),
        ("ERP-Gründerkredit - StartGeld", "200.000 EUR", "120.000 Investition + maximal 80.000 Betriebsmittel"),
        ("zusätzliche Eigenmittel im Base-Case", "0 EUR", "bestehendes gesetzliches Stammkapital separat nachweisen"),
        ("Finanzierungssumme", "200.000 EUR", "entspricht dem verbindlichen Base-Case"),
    ]
    make_table(
        doc,
        ["Finanzierung", "Betrag", "Erläuterung"],
        funding_rows,
        widths=[8.2, 3.1, 6.7],
        font_size=7.8,
        total_rows={1, 4},
    )

    doc.add_paragraph()
    add_note(
        doc,
        "Entscheidende Förderbedingung",
        "Die Hausbank muss vor Beauftragung bestätigen, dass das 110.000-EUR-Software-/IP-Paket und die 10.000-EUR-Hardware "
        "als mittel-/langfristige Investitionen anerkannt werden. Wird dies abgelehnt, ist der 200.000-EUR-" 
        "Antrag in dieser Form nicht tragfähig; dann den Kredit reduzieren oder die "
        "Investition separat mit Eigenkapital finanzieren.",
        fill=WARM,
    )
    doc.add_paragraph()
    add_note(
        doc,
        "Abgrenzung zum Jahresbudget",
        "Das 36-Monats-Modell verknüpft Investitionsmeilensteine, 0 EUR verbindliche Jahr-1-Personalkosten, fixe Kosten, "
        "variable COGS, monatliche Kundenzahlungen und Kapitaldienst. Ohne zusätzliches Eigenkapital beträgt die "
        "Endliquidität Jahr 1 75,640 TEUR; der Cash-Tiefpunkt liegt mit 41,046 TEUR in Jahr 2. Der optionale Personalrahmen "
        "bis 80 TEUR ist darin nicht enthalten und benötigt vor Freigabe mindestens rund 68,95 TEUR zusätzliche Liquidität, "
        "wenn der 30-TEUR-Puffer über 36 Monate gehalten werden soll.",
    )
    doc.add_paragraph()
    add_note(
        doc,
        "Quellenbasis",
        "KfW-Formular 600 000 5282; KfW-Merkblatt 067, Stand 18.06.2026; KfW-Produktseite abgerufen 12.07.2026; "
        "Consultry Businessplan, Product Vision v2.7, Alignment Control Plane v1.1 und MVP-PRD. Die modulare Entwicklungsdarstellung ist von der Struktur der OmniSEC-Planung inspiriert; "
        "Umfang und Beträge sind ausschließlich Consultry-Planannahmen und durch Angebote zu ersetzen.",
    )

    path = OUT_DOC / "04_Consultry_Investitionsplan.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    OUT_DOC.mkdir(parents=True, exist_ok=True)
    modes = set(sys.argv[1:] or ["all"])
    if "all" in modes or "rentability" in modes:
        print(create_rentability())
    if "all" in modes or "liquidity" in modes:
        print(create_liquidity())
    if "all" in modes or "investment" in modes:
        print(create_investment())
