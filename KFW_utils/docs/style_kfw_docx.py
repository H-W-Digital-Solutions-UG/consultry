from pathlib import Path
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


GREEN = "0F5B4D"
DARK = "17332D"
MINT = "DCEBE6"
PALE = "F4F7F5"
WARM = "FFF0DB"
GRAY = "59645F"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant = OxmlElement("w:cantSplit")
    tr_pr.append(cant)


def set_cell_text(cell, color=None, bold=None, size=8.5):
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        for run in paragraph.runs:
            run.font.name = "Aptos"
            run.font.size = Pt(size)
            if color:
                run.font.color.rgb = RGBColor.from_string(color)
            if bold is not None:
                run.bold = bold


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])


def add_bottom_border(paragraph, color=GREEN, size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def style_doc(path):
    doc = Document(path)
    doc.core_properties.title = "Consultry Geschäftskonzept / Businessplan"
    doc.core_properties.subject = "ERP-Gründerkredit StartGeld 067"
    doc.core_properties.author = "Consultry Software GmbH"

    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.7)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(1.8)
        section.header_distance = Cm(0.7)
        section.footer_distance = Cm(0.7)

        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        try:
            hp.add_run().add_picture(
                str(Path("design/logos/consultry-weblogo-cropped.png").resolve()),
                width=Inches(1.25),
            )
        except Exception:
            run = hp.add_run("CONSULTRY")
            run.bold = True
            run.font.color.rgb = RGBColor.from_string(GREEN)
            run.font.size = Pt(11)
        run = hp.add_run("    Geschäftskonzept | KfW StartGeld 067 | ENTWURF")
        run.font.name = "Aptos"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(GRAY)
        add_bottom_border(hp, GREEN, "5")

        footer = section.footer
        fp = footer.paragraphs[0]
        fp.add_run("Consultry Software GmbH | Planungsstand 12.07.2026 | Seite ")
        for run in fp.runs:
            run.font.name = "Aptos"
            run.font.size = Pt(7.5)
            run.font.color.rgb = RGBColor.from_string(GRAY)
        add_page_field(fp)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.4)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08
    normal.paragraph_format.widow_control = True

    for name in ("Body Text", "First Paragraph"):
        st = doc.styles[name]
        st.font.name = "Aptos"
        st.font.size = Pt(9.4)
        st.font.color.rgb = RGBColor.from_string(DARK)
        st.paragraph_format.space_after = Pt(5)
        st.paragraph_format.line_spacing = 1.08

    title = doc.styles["Title"]
    title.font.name = "Aptos Display"
    title.font.size = Pt(32)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(GREEN)
    title.paragraph_format.space_before = Pt(80)
    title.paragraph_format.space_after = Pt(10)

    subtitle = doc.styles["Subtitle"]
    subtitle.font.name = "Aptos Display"
    subtitle.font.size = Pt(20)
    subtitle.font.color.rgb = RGBColor.from_string(DARK)
    subtitle.paragraph_format.space_after = Pt(20)

    date = doc.styles["Date"]
    date.font.name = "Aptos"
    date.font.size = Pt(11)
    date.font.color.rgb = RGBColor.from_string(GRAY)
    date.paragraph_format.space_after = Pt(30)

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Aptos Display"
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor.from_string(GREEN)
    h1.paragraph_format.space_before = Pt(4)
    h1.paragraph_format.space_after = Pt(8)
    h1.paragraph_format.keep_with_next = True
    h1.paragraph_format.page_break_before = True

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Aptos Display"
    h2.font.size = Pt(12.5)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor.from_string(DARK)
    h2.paragraph_format.space_before = Pt(10)
    h2.paragraph_format.space_after = Pt(4)
    h2.paragraph_format.keep_with_next = True

    if "Compact" in doc.styles:
        compact = doc.styles["Compact"]
        compact.font.name = "Aptos"
        compact.font.size = Pt(9.1)
        compact.font.color.rgb = RGBColor.from_string(DARK)
        compact.paragraph_format.space_after = Pt(2.5)
        compact.paragraph_format.line_spacing = 1.03
        compact.paragraph_format.left_indent = Cm(0.25)

    if "Block Text" in doc.styles:
        block = doc.styles["Block Text"]
        block.font.name = "Aptos"
        block.font.size = Pt(9.5)
        block.font.bold = True
        block.font.color.rgb = RGBColor.from_string(DARK)
        block.paragraph_format.left_indent = Cm(0.5)
        block.paragraph_format.right_indent = Cm(0.5)
        block.paragraph_format.space_before = Pt(20)
        block.paragraph_format.space_after = Pt(20)

    # Ensure a true cover page. The TOC follows on its own page, and every H1 starts a new page.
    if len(doc.paragraphs) > 3:
        doc.paragraphs[3].add_run().add_break(WD_BREAK.PAGE)

    for idx, table in enumerate(doc.tables):
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        if table.rows:
            set_repeat_table_header(table.rows[0])
        for r_i, row in enumerate(table.rows):
            set_cant_split(row)
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                set_cell_margins(cell)
                if r_i == 0:
                    set_cell_shading(cell, GREEN)
                    set_cell_text(cell, "FFFFFF", True, 8.2)
                else:
                    if r_i % 2 == 0:
                        set_cell_shading(cell, PALE)
                    set_cell_text(cell, DARK, None, 8.2)

    # Cover styling and warning emphasis.
    for p in doc.paragraphs[:5]:
        if p.style.name in ("Title", "Subtitle", "Date"):
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if len(doc.paragraphs) > 3:
        p = doc.paragraphs[3]
        p_pr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), WARM)
        p_pr.append(shd)

    # Keep the compact capital-service subsection together on its own page.
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith("10.3 Kreditkonditionen und Kapitaldienst"):
            paragraph.paragraph_format.page_break_before = True
            break

    doc.save(path)


if __name__ == "__main__":
    style_doc(sys.argv[1])
